import os
import re

from docx import Document
from ordered_set import OrderedSet

from funcTrace.GetFig import getFig
from funcTrace.RiskFuncManage import FunctionManager


def extract_custom_headers_and_sources(header_file, included_files=None, included_sources=None):
    if included_files is None:
        included_files = set()
    if included_sources is None:
        included_sources = set()

    # 用于匹配非官方库的正则表达式
    custom_header_pattern = re.compile(r'#include\s+"(.+?)"')

    with open(header_file, 'r') as file:
        content = file.read()

    # 查找所有匹配的非官方库文件
    custom_headers = custom_header_pattern.findall(content)

    for header in custom_headers:
        # 如果该文件尚未处理，递归解析该文件
        if header not in included_files:
            header_path = os.path.join(os.path.dirname(header_file), header)
            included_files.add(header_path)
            if os.path.exists(header_path):
                extract_custom_headers_and_sources(header_path, included_files, included_sources)
            else:
                print(f"Warning: {header_path} does not exist.")

            # 尝试找到对应的实现文件
            source_file = os.path.splitext(header_path)[0] + '.c'
            if os.path.exists(source_file):
                included_sources.add(source_file)

    return included_files, included_sources


def export_report_to_txt(file_name: str, custom_sources):
    report_set = OrderedSet()
    for path in custom_sources:
        manager = FunctionManager(path)
        report_set.add(manager.riskFunction())

        if len(manager.get_fig_sizes_high()) != 0:
            getFig(manager.get_fig_labels_high(), manager.get_fig_sizes_high(), os.path.dirname(file_name),
                   "High risk",
                   f"{os.path.basename(path)}_high_risk.jpg", True).get_fig()
        if len(manager.get_fig_sizes_medium()) != 0:
            getFig(manager.get_fig_labels_medium(), manager.get_fig_sizes_medium(), os.path.dirname(file_name),
                   "Medium risk",
                   f"{os.path.basename(path)}_medium_risk.jpg", True).get_fig()
        if len(manager.get_fig_sizes_low()) != 0:
            getFig(manager.get_fig_labels_low(), manager.get_fig_sizes_low(), os.path.dirname(file_name),
                   "Low risk",
                   f"{os.path.basename(path)}_low_risk.jpg", True).get_fig()

    text = ""
    for report in report_set:
        text = text + report + "\n\n\n"

    with open(file_name, 'w') as f:
        f.write(text)


def export_report_to_doc(file_name: str, custom_sources):
    doc = Document()
    for path in custom_sources:
        manager = FunctionManager(path)
        text = manager.riskFunction()

        doc.add_paragraph(text)

        if len(manager.get_fig_sizes_high()) != 0:
            high = getFig(manager.get_fig_labels_high(), manager.get_fig_sizes_high(), os.path.dirname(file_name),
                          "High risk",
                          f"{os.path.basename(path)}_high_risk.jpg", True).get_fig()
            doc.add_picture(high)
            os.remove(high)
        if len(manager.get_fig_sizes_medium()) != 0:
            medium = getFig(manager.get_fig_labels_medium(), manager.get_fig_sizes_medium(), os.path.dirname(file_name),
                            "Medium risk",
                            f"{os.path.basename(path)}_medium_risk.jpg", True).get_fig()
            doc.add_picture(medium)
            os.remove(medium)
        if len(manager.get_fig_sizes_low()) != 0:
            low = getFig(manager.get_fig_labels_low(), manager.get_fig_sizes_low(), os.path.dirname(file_name),
                         "Low risk",
                         f"{os.path.basename(path)}_low_risk.jpg", True).get_fig()
            doc.add_picture(low)
            os.remove(low)

        doc.add_page_break()

    doc.save(file_name)
